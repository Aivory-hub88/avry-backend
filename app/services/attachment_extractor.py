"""
Attachment extraction for chat agents.

Customers send documents and images to their agents over Telegram/Slack.
This module turns raw file bytes into text the LLM can read:

  - Documents (PDF/Word/Excel/CSV/plain text) -> extracted text
  - Images (photo/scan) -> a description + any legible text, via a vision model

Everything here treats file bytes as UNTRUSTED input: we only ever parse/extract,
never execute, and we cap both the input size and the extracted text length so a
huge or malicious file can't blow up the prompt or the process.
"""

import csv
import io
import logging
import os
from typing import Optional

import requests

logger = logging.getLogger(__name__)

# Hard limits — a document past these is summarised as "too large" rather than parsed
MAX_FILE_BYTES = 20 * 1024 * 1024          # 20 MB (matches Telegram Bot API download cap)
MAX_EXTRACTED_CHARS = 12000                 # keep the injected context bounded
MAX_IMAGE_BYTES = 8 * 1024 * 1024           # vision payloads stay small

# Vision model on OpenRouter (deepseek-v4-flash is text-only, so images need this)
VISION_MODEL = os.getenv("VISION_MODEL", "google/gemini-2.5-flash")
OPENROUTER_URL = "https://openrouter.ai/api/v1/chat/completions"

_TEXT_EXTS = {".txt", ".md", ".log", ".json", ".yaml", ".yml", ".text"}
_IMAGE_EXTS = {".jpg", ".jpeg", ".png", ".webp", ".gif", ".bmp"}


def _truncate(text: str) -> str:
    text = (text or "").strip()
    if len(text) > MAX_EXTRACTED_CHARS:
        return text[:MAX_EXTRACTED_CHARS] + "\n\n[...document truncated...]"
    return text


def _ext(filename: str) -> str:
    filename = (filename or "").lower()
    dot = filename.rfind(".")
    return filename[dot:] if dot != -1 else ""


def is_image(filename: str, mime: Optional[str] = None) -> bool:
    if mime and mime.startswith("image/"):
        return True
    return _ext(filename) in _IMAGE_EXTS


# ============================================================================
# DOCUMENT EXTRACTION
# ============================================================================

def _extract_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    parts = []
    for page in reader.pages:
        parts.append(page.extract_text() or "")
        if sum(len(p) for p in parts) > MAX_EXTRACTED_CHARS:
            break
    return "\n".join(parts)


def _extract_docx(data: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(data))
    parts = [p.text for p in doc.paragraphs if p.text]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def _extract_xlsx(data: bytes) -> str:
    from openpyxl import load_workbook

    wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    parts = []
    for ws in wb.worksheets:
        parts.append(f"# Sheet: {ws.title}")
        for row in ws.iter_rows(values_only=True):
            cells = ["" if c is None else str(c) for c in row]
            if any(cells):
                parts.append(" | ".join(cells))
            if sum(len(p) for p in parts) > MAX_EXTRACTED_CHARS:
                break
    wb.close()
    return "\n".join(parts)


def _extract_csv(data: bytes) -> str:
    text = data.decode("utf-8", errors="replace")
    rows = list(csv.reader(io.StringIO(text)))
    return "\n".join(" | ".join(r) for r in rows)


def extract_document_text(filename: str, data: bytes, mime: Optional[str] = None) -> Optional[str]:
    """Extract readable text from a document. Returns None if unsupported."""
    if len(data) > MAX_FILE_BYTES:
        return "[This document is too large to read (over 20 MB).]"

    ext = _ext(filename)
    try:
        if ext == ".pdf" or mime == "application/pdf":
            return _truncate(_extract_pdf(data))
        if ext == ".docx" or (mime or "").endswith("wordprocessingml.document"):
            return _truncate(_extract_docx(data))
        if ext in (".xlsx", ".xlsm") or (mime or "").endswith("spreadsheetml.sheet"):
            return _truncate(_extract_xlsx(data))
        if ext == ".csv" or mime == "text/csv":
            return _truncate(_extract_csv(data))
        if ext in _TEXT_EXTS or (mime or "").startswith("text/"):
            return _truncate(data.decode("utf-8", errors="replace"))
    except Exception as e:
        logger.error(f"Document extraction failed for {filename}: {e}")
        return f"[Could not read '{filename}' — the file may be corrupt or password-protected.]"
    return None


# ============================================================================
# IMAGE UNDERSTANDING (vision)
# ============================================================================

def describe_image(data: bytes, mime: str, user_caption: str = "") -> Optional[str]:
    """Send an image to a vision model and return a description + any legible text."""
    api_key = os.getenv("OPENROUTER_API_KEY")
    if not api_key:
        logger.warning("OPENROUTER_API_KEY not set; cannot process image")
        return None
    if len(data) > MAX_IMAGE_BYTES:
        return "[This image is too large to analyse.]"

    import base64

    b64 = base64.b64encode(data).decode()
    prompt = (
        "Describe this image for a business assistant. If it contains a document, "
        "invoice, receipt, form, or screenshot, transcribe ALL text you can read, "
        "preserving numbers and structure. Be concise but complete."
    )
    if user_caption:
        prompt += f'\nThe user said: "{user_caption}"'

    try:
        resp = requests.post(
            OPENROUTER_URL,
            headers={
                "Authorization": f"Bearer {api_key}",
                "Content-Type": "application/json",
                "HTTP-Referer": "https://aivory.id",
                "X-Title": "Aivory Agent Vision",
            },
            json={
                "model": VISION_MODEL,
                "messages": [
                    {
                        "role": "user",
                        "content": [
                            {"type": "text", "text": prompt},
                            {"type": "image_url", "image_url": {"url": f"data:{mime};base64,{b64}"}},
                        ],
                    }
                ],
            },
            timeout=60,
        )
        if resp.ok:
            content = resp.json().get("choices", [{}])[0].get("message", {}).get("content")
            return _truncate(content) if content else None
        logger.error(f"Vision model {resp.status_code}: {resp.text[:200]}")
    except (requests.RequestException, ValueError) as e:
        logger.error(f"Vision request failed: {e}")
    return None


# ============================================================================
# COMPOSE — build the final prompt text handed to the agent gateway
# ============================================================================

def compose_prompt(user_text: str, attachments: list) -> str:
    """Combine the user's typed text with extracted attachment content.

    attachments: list of {"filename": str, "content": str, "kind": "document"|"image"}
    """
    blocks = []
    for att in attachments:
        content = (att.get("content") or "").strip()
        if not content:
            continue
        label = "Image" if att.get("kind") == "image" else "Document"
        name = att.get("filename") or "attachment"
        blocks.append(f"[{label}: {name}]\n{content}")

    user_text = (user_text or "").strip()
    if not blocks:
        return user_text

    joined = "\n\n".join(blocks)
    if user_text:
        return f"{joined}\n\n[User's message]\n{user_text}"
    # No caption: tell the agent the user sent the file with no words
    return f"{joined}\n\n[The user sent this without any message. Respond helpfully based on its content.]"
